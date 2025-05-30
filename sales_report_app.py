import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import os
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO

# Page configuration
st.set_page_config(
    page_title="Weekly Sales Report Generator",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #ff6b6b;
    }
    .success-metric {
        border-left-color: #51cf66;
    }
    .warning-metric {
        border-left-color: #ffd43b;
    }
    .danger-metric {
        border-left-color: #ff6b6b;
    }
</style>
""", unsafe_allow_html=True)

st.title("📊 Weekly Sales Report Generator")
st.markdown("Generate comprehensive weekly sales reports with automated calculations and insights.")

# Sidebar for navigation
with st.sidebar:
    st.header("📋 Navigation")
    selected_section = st.radio(
        "Select Section:",
        ["Data Input", "Analytics Dashboard", "Report Preview", "Generate Report"]
    )

# Initialize session state
if 'report_data' not in st.session_state:
    st.session_state.report_data = {}

# Helper functions
def format_currency(amount):
    """Format currency with KSH prefix and commas"""
    return f"KSH {amount:,.0f}"

def calculate_percentage(numerator, denominator):
    """Calculate percentage with error handling"""
    if denominator == 0:
        return 0
    return (numerator / denominator) * 100

def create_performance_chart(data):
    """Create performance visualization"""
    fig = go.Figure()
    
    fig.add_trace(go.Bar(
        name='Budget',
        x=['Weekly', 'MTD'],
        y=[data['weekly_budget'], data['budget']],
        marker_color='lightblue'
    ))
    
    fig.add_trace(go.Bar(
        name='Actual',
        x=['Weekly', 'MTD'],
        y=[data['current_week_revenue'], data['mtd_revenue']],
        marker_color='darkblue'
    ))
    
    fig.update_layout(
        title='Budget vs Actual Performance',
        xaxis_title='Period',
        yaxis_title='Amount (KSH)',
        barmode='group',
        height=400
    )
    
    return fig

# Section 1: Data Input
if selected_section == "Data Input":
    st.header("🔹 1. Performance Data Input")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Monthly Targets & Achievement")
        budget = st.number_input("Monthly Budget (KSH)", value=113998325, format="%d")
        mtd_revenue = st.number_input("MTD Revenue (KSH)", value=93415418, format="%d")
        
        st.subheader("Projection Estimates")
        historical_trend = st.number_input("Historical Trend (KSH)", value=89936015, format="%d")
        linear_extrap = st.number_input("Linear Extrapolation (KSH)", value=99857861, format="%d")
        blended_estimate = st.number_input("Blended Conservative Estimate (KSH)", value=93904753, format="%d")
    
    with col2:
        st.subheader("Weekly Performance")
        weekly_budget = st.number_input("Weekly Budget (KSH)", value=26479125, format="%d")
        current_week_revenue = st.number_input("Current Week Revenue (KSH)", value=20943811, format="%d")
        previous_week_revenue = st.number_input("Previous Week Revenue (KSH)", value=20353938, format="%d")
        
        st.subheader("Operational Metrics")
        short_supplies = st.number_input("Short Supplies (KSH)", value=1266460, format="%d")
        returns = st.number_input("Returns (KSH)", value=193615, format="%d")
    
    st.header("🔹 2. Additional Information")
    
    col3, col4 = st.columns(2)
    with col3:
        week_number = st.number_input("Week Number", value=22, min_value=1, max_value=52)
        report_date = st.date_input("Report Date", value=datetime.now().date())
    
    with col4:
        highlight_may_25 = st.checkbox("✅ May 25 sales exceeded trend", value=True)
        parmesan_price_increase = st.checkbox("✅ Due to Parmesan price increase", value=True)
    
    st.header("🔹 3. Upload Supplementary Data")
    col5, col6 = st.columns(2)
    
    with col5:
        short_supply_file = st.file_uploader("Upload Top 10 Short Supplied Items (Excel)", type=["xlsx", "csv"])
    
    with col6:
        market_return_file = st.file_uploader("Upload Top 10 Market Returns (Excel)", type=["xlsx", "csv"])
    
    # Store data in session state
    st.session_state.report_data = {
        'budget': budget,
        'mtd_revenue': mtd_revenue,
        'weekly_budget': weekly_budget,
        'current_week_revenue': current_week_revenue,
        'previous_week_revenue': previous_week_revenue,
        'short_supplies': short_supplies,
        'returns': returns,
        'historical_trend': historical_trend,
        'linear_extrap': linear_extrap,
        'blended_estimate': blended_estimate,
        'week_number': week_number,
        'report_date': report_date,
        'highlight_may_25': highlight_may_25,
        'parmesan_price_increase': parmesan_price_increase,
        'short_supply_file': short_supply_file,
        'market_return_file': market_return_file
    }
    
    if st.button("💾 Save Data & Continue"):
        st.success("Data saved successfully! Navigate to other sections using the sidebar.")

# Section 2: Analytics Dashboard
elif selected_section == "Analytics Dashboard":
    st.header("📈 Analytics Dashboard")
    
    if not st.session_state.report_data:
        st.warning("Please input data in the 'Data Input' section first.")
    else:
        data = st.session_state.report_data
        
        # Calculate key metrics
        revenue_gap = data['budget'] - data['mtd_revenue']
        achievement_pct = calculate_percentage(data['mtd_revenue'], data['budget'])
        weekly_variance = data['current_week_revenue'] - data['weekly_budget']
        growth_rate = calculate_percentage(
            data['current_week_revenue'] - data['previous_week_revenue'], 
            data['previous_week_revenue']
        )
        closing_pct = calculate_percentage(data['blended_estimate'], data['budget'])
        
        # Key Metrics Display
        st.subheader("📊 Key Performance Indicators")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric(
                label="MTD Achievement",
                value=f"{achievement_pct:.1f}%",
                delta=f"{achievement_pct - 100:.1f}% vs Budget"
            )
        
        with col2:
            st.metric(
                label="Weekly Growth",
                value=f"{growth_rate:+.2f}%",
                delta="Week-over-Week"
            )
        
        with col3:
            st.metric(
                label="Revenue Gap",
                value=format_currency(revenue_gap),
                delta=f"{100 - achievement_pct:.1f}% remaining"
            )
        
        with col4:
            st.metric(
                label="Projected Closing",
                value=f"{closing_pct:.1f}%",
                delta="of Budget"
            )
        
        # Visualizations
        st.subheader("📊 Performance Visualizations")
        
        col5, col6 = st.columns(2)
        
        with col5:
            # Budget vs Actual Chart
            chart = create_performance_chart(data)
            st.plotly_chart(chart, use_container_width=True)
        
        with col6:
            # Projection Scenarios
            scenarios_data = {
                'Scenario': ['Historical Trend', 'Linear Extrapolation', 'Blended Estimate'],
                'Amount': [data['historical_trend'], data['linear_extrap'], data['blended_estimate']],
                'Achievement %': [
                    calculate_percentage(data['historical_trend'], data['budget']),
                    calculate_percentage(data['linear_extrap'], data['budget']),
                    calculate_percentage(data['blended_estimate'], data['budget'])
                ]
            }
            
            fig_scenarios = px.bar(
                x=scenarios_data['Scenario'], 
                y=scenarios_data['Achievement %'],
                title='Closing Scenarios (% of Budget)',
                labels={'x': 'Scenario', 'y': 'Achievement %'}
            )
            fig_scenarios.add_hline(y=100, line_dash="dash", line_color="red", 
                                  annotation_text="Budget Target")
            st.plotly_chart(fig_scenarios, use_container_width=True)

# Section 3: Report Preview
elif selected_section == "Report Preview":
    st.header("👁️ Report Preview")
    
    if not st.session_state.report_data:
        st.warning("Please input data in the 'Data Input' section first.")
    else:
        data = st.session_state.report_data
        
        # Calculate metrics for preview
        revenue_gap = data['budget'] - data['mtd_revenue']
        achievement_pct = calculate_percentage(data['mtd_revenue'], data['budget'])
        weekly_variance = data['current_week_revenue'] - data['weekly_budget']
        growth_rate = calculate_percentage(
            data['current_week_revenue'] - data['previous_week_revenue'], 
            data['previous_week_revenue']
        )
        closing_pct = calculate_percentage(data['blended_estimate'], data['budget'])
        
        st.markdown(f"# Week {data['week_number']} Sales Report")
        st.markdown(f"**Generated on:** {data['report_date']}")
        
        st.markdown("## MTD Sales Revenue Update")
        st.write(f"**Budget:** {format_currency(data['budget'])}")
        st.write(f"**MTD Revenue:** {format_currency(data['mtd_revenue'])}")
        st.write(f"**Achievement vs Budget:** {achievement_pct:.0f}%")
        st.write(f"**Revenue Gap to Budget:** {format_currency(revenue_gap)} ({100 - achievement_pct:.0f}%)")
        st.write(f"**Closing Revenue Estimate:** {format_currency(data['blended_estimate'])} ({closing_pct:.0f}% of Budget)")
        
        st.markdown("## Current Week Performance")
        st.write(f"**Weekly Budget:** {format_currency(data['weekly_budget'])}")
        st.write(f"**Current Week Revenue:** {format_currency(data['current_week_revenue'])}")
        st.write(f"**Variance to Budget:** {format_currency(weekly_variance)} ({(weekly_variance / data['weekly_budget'] * 100):+.0f}%)")
        st.write(f"**Previous Week Revenue:** {format_currency(data['previous_week_revenue'])}")
        st.write(f"**Growth Rate:** {growth_rate:.2f}%")
        
        st.markdown("## Operational Insights")
        st.write(f"**Short Supplies:** {format_currency(data['short_supplies'])} (~{(data['short_supplies'] / data['current_week_revenue'] * 100):.1f}% impact)")
        st.write(f"**Returns:** {format_currency(data['returns'])} (~{(data['returns'] / data['current_week_revenue'] * 100):.1f}% impact)")
        
        st.markdown("## Key Highlights")
        st.write(f"Week {data['week_number']} showed a {growth_rate:+.2f}% growth over Week {data['week_number']-1}, though still under budget.")
        st.write(f"MTD Revenue is now at {achievement_pct:.0f}% of budget with {format_currency(revenue_gap)} to go.")
        
        if data['highlight_may_25']:
            st.write("• May 25 sales exceeded the historical trend.")
            if data['parmesan_price_increase']:
                st.write("• This was likely due to an increase in Parmesan price.")

# Section 4: Generate Report
elif selected_section == "Generate Report":
    st.header("📝 Generate Report")
    
    if not st.session_state.report_data:
        st.warning("Please input data in the 'Data Input' section first.")
    else:
        data = st.session_state.report_data
        
        st.info("Review your data and click 'Generate Report' to create the Word document.")
        
        # Quick summary
        col1, col2 = st.columns(2)
        with col1:
            st.write("**Week Number:**", data['week_number'])
            st.write("**Report Date:**", data['report_date'])
            st.write("**MTD Revenue:**", format_currency(data['mtd_revenue']))
        
        with col2:
            achievement_pct = calculate_percentage(data['mtd_revenue'], data['budget'])
            st.write("**Monthly Budget:**", format_currency(data['budget']))
            st.write("**Achievement:**", f"{achievement_pct:.1f}%")
            st.write("**Projected Closing:**", format_currency(data['blended_estimate']))
        
        if st.button("📝 Generate Word Report", type="primary"):
            try:
                # Calculate all metrics
                revenue_gap = data['budget'] - data['mtd_revenue']
                achievement_pct = calculate_percentage(data['mtd_revenue'], data['budget'])
                weekly_variance = data['current_week_revenue'] - data['weekly_budget']
                growth_rate = calculate_percentage(
                    data['current_week_revenue'] - data['previous_week_revenue'], 
                    data['previous_week_revenue']
                )
                closing_pct = calculate_percentage(data['blended_estimate'], data['budget'])
                
                # Create Word document
                doc = Document()
                
                # Title and date
                title = doc.add_heading(f"Week {data['week_number']} Sales Report", 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                date_para = doc.add_paragraph(f"Generated on: {data['report_date']}")
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # MTD Sales Revenue Update
                doc.add_heading("MTD Sales Revenue Update", level=1)
                doc.add_paragraph(f"Budget: {format_currency(data['budget'])}")
                doc.add_paragraph(f"MTD Revenue: {format_currency(data['mtd_revenue'])}")
                doc.add_paragraph(f"Achievement vs Budget: {achievement_pct:.0f}%")
                doc.add_paragraph(f"Revenue Gap to Budget: {format_currency(revenue_gap)} ({100 - achievement_pct:.0f}%)")
                doc.add_paragraph(f"Closing Revenue Estimate: {format_currency(data['blended_estimate'])} ({closing_pct:.0f}% of Budget)")
                
                # Current Week Performance
                doc.add_heading("Current Week Performance", level=1)
                doc.add_paragraph(f"Weekly Budget: {format_currency(data['weekly_budget'])}")
                doc.add_paragraph(f"Current Week Revenue: {format_currency(data['current_week_revenue'])}")
                doc.add_paragraph(f"Variance to Budget: {format_currency(weekly_variance)} ({(weekly_variance / data['weekly_budget'] * 100):+.0f}%)")
                doc.add_paragraph(f"Previous Week Revenue: {format_currency(data['previous_week_revenue'])}")
                doc.add_paragraph(f"Growth Rate: {growth_rate:.2f}%")
                
                # Operational Insights
                doc.add_heading("Operational Insights", level=1)
                doc.add_paragraph(f"Short Supplies: {format_currency(data['short_supplies'])} (~{(data['short_supplies'] / data['current_week_revenue'] * 100):.1f}% impact)")
                doc.add_paragraph(f"Returns: {format_currency(data['returns'])} (~{(data['returns'] / data['current_week_revenue'] * 100):.1f}% impact)")
                
                # Key Highlights
                doc.add_heading("Key Highlights", level=1)
                doc.add_paragraph(f"Week {data['week_number']} showed a {growth_rate:+.2f}% growth over Week {data['week_number']-1}, though still under budget.")
                doc.add_paragraph(f"MTD Revenue is now at {achievement_pct:.0f}% of budget with {format_currency(revenue_gap)} to go.")
                
                if data['highlight_may_25']:
                    doc.add_paragraph("• May 25 sales exceeded the historical trend.", style="List Bullet")
                    if data['parmesan_price_increase']:
                        doc.add_paragraph("• This was likely due to an increase in Parmesan price.", style="List Bullet")
                
                # Closing Estimates Summary
                doc.add_heading("Closing Estimates Summary", level=1)
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Scenario'
                hdr_cells[1].text = 'Estimate (KSH)'
                hdr_cells[2].text = '% of Budget'
                
                scenarios = [
                    ('Historical Trend', data['historical_trend']),
                    ('Linear Extrapolation', data['linear_extrap']),
                    ('Blended Estimate', data['blended_estimate'])
                ]
                
                for label, value in scenarios:
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = f"{value:,.0f}"
                    row_cells[2].text = f"{(value / data['budget'] * 100):.1f}%"
                
                # Add supplementary data if uploaded
                if data['short_supply_file']:
                    doc.add_heading("Top 10 Short Supplied Items", level=1)
                    try:
                        if data['short_supply_file'].name.endswith('.csv'):
                            df_short = pd.read_csv(data['short_supply_file'])
                        else:
                            df_short = pd.read_excel(data['short_supply_file'])
                        doc.add_paragraph(df_short.to_string(index=False))
                    except Exception as e:
                        doc.add_paragraph(f"Error reading short supply file: {str(e)}")
                
                if data['market_return_file']:
                    doc.add_heading("Top 10 Market Returns", level=1)
                    try:
                        if data['market_return_file'].name.endswith('.csv'):
                            df_returns = pd.read_csv(data['market_return_file'])
                        else:
                            df_returns = pd.read_excel(data['market_return_file'])
                        doc.add_paragraph(df_returns.to_string(index=False))
                    except Exception as e:
                        doc.add_paragraph(f"Error reading market returns file: {str(e)}")
                
                # Save to BytesIO buffer
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                # Generate filename
                filename = f"Week{data['week_number']}_Sales_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                
                # Download button
                st.download_button(
                    label="📥 Download Word Report",
                    data=buffer.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("✅ Report generated successfully!")
                
            except Exception as e:
                st.error(f"❌ Error generating report: {str(e)}")

# Footer
st.markdown("---")
st.markdown("*Weekly Sales Report Generator - Built with Streamlit*")
   
